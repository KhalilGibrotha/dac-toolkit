"""
sections.py — Section break management and body page footer.

Section structure produced by the builder:
  Section 1: Cover page    — page border, no footer, watermark if Draft
  Section 2: TOC           — no border, page numbering restarted at 1
  Section 3: Revision      — no border, linked header to previous = False
  Section 4: Body content  — no border, "page N" footer right-aligned

Each section sets is_linked_to_previous = False on its header to prevent
header content (especially the watermark) from bleeding across sections.
"""

from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .constants import GRAY_TEXT, FONT_BODY
from .xml_helpers import set_run_color, para_spacing


def add_section_break(doc, break_type=WD_SECTION.NEW_PAGE):
    """
    Add a section break and return the new section object.
    Use the returned section to set margins and header/footer linkage
    before adding any content to the new section.
    """
    return doc.add_section(break_type)


def add_body_footer(section, meta: dict):
    """
    Add a right-aligned 'page N' footer to the body section.

    The page number is inserted as a Word PAGE field (w:instrText) so that
    Word renders the actual page number. A literal 'page ' prefix precedes
    the field, matching the reference document style.

    Args:
        section: the python-docx section to apply the footer to
        meta:    front matter dict (reserved for future use, e.g. doc title in footer)
    """
    footer = section.footer
    # Remove any inherited paragraphs
    for para in footer.paragraphs:
        p = para._p
        try:
            p.getparent().remove(p)
        except Exception:
            pass

    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para_spacing(p, before=0, after=0)

    # "page " literal text
    run1 = p.add_run("page ")
    run1.font.name = FONT_BODY
    run1.font.size = Pt(9)
    set_run_color(run1, GRAY_TEXT)

    # PAGE field: begin → instrText → end
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run2 = p.add_run()
    run2._r.append(fldChar1)
    run3 = p.add_run()
    run3._r.append(instrText)
    run4 = p.add_run()
    run4._r.append(fldChar2)

    for r in (run2, run3, run4):
        r.font.name = FONT_BODY
        r.font.size = Pt(9)
        set_run_color(r, GRAY_TEXT)
