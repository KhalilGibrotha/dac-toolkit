"""
cover_page.py — Cover page and draft watermark rendering.

FRAGILITY WARNING
-----------------
This module contains two pieces of code that are unusually sensitive to
changes. Read the notes before modifying either.

1. add_draft_watermark()
   Injects a VML-based DRAFT watermark into the section header as a raw XML
   string. VML (Vector Markup Language) is the legacy Office drawing format
   and is the most broadly compatible way to render Word watermarks. The
   namespace declarations are embedded in the element string because
   etree.fromstring() builds an isolated element tree — they cannot be
   inherited from a parent. Do not restructure this XML string. The rotation
   value (315 = -45 degrees), size, and z-index are all intentional.

2. build_cover_page()
   The cover layout uses counted empty paragraphs (spacers) to vertically
   position the title block and copyright block. This is not elegant but it
   is reliable and exactly matches the reference document. Replacing spacers
   with text box positioning or table-based layout will change the visual
   output. If the number of spacer paragraphs needs tuning, adjust the
   loop counts at the "Vertical spacer" and "Spacer to push copyright block"
   comments.
"""

import os
from lxml import etree

from docx.shared import Inches, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .constants import (
    BLUE_DARK, BLUE_LINK, BLACK, DARK_NAVY,
    FONT_BODY, FONT_TITLE,
    ORG_NAME, ORG_DEPT, ORG_ADDR1, ORG_ADDR2, ORG_URL,
    PAGE_W_TWIPS, PAGE_H_TWIPS, TWIPS_TO_EMU,
)
from .xml_helpers import (
    set_run_color, para_spacing, set_paragraph_border_bottom,
)


# ── Draft watermark ───────────────────────────────────────────────────────────

def add_draft_watermark(section):
    """
    Inject a diagonal 'DRAFT' text watermark into the section header via VML.

    The watermark renders on every page in the section because it lives in
    the section header. The VML approach is used (rather than DrawingML) for
    maximum compatibility with both Word and LibreOffice.

    DO NOT modify the XML string without testing in both applications.
    namespace declarations must remain inline (not inherited from parent).
    """
    header = section.header
    # Remove existing header content
    for para in header.paragraphs:
        p = para._p
        p.getparent().remove(p)

    hdr_xml = header._element

    wmark_xml = (
        '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        '     xmlns:v="urn:schemas-microsoft-com:vml"'
        '     xmlns:o="urn:schemas-microsoft-com:office:office"'
        '     xmlns:w10="urn:schemas-microsoft-com:office:word">'
        '  <w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>'
        '  <w:r>'
        '    <w:rPr><w:noProof/></w:rPr>'
        '    <w:pict>'
        '      <v:shape id="watermark" type="#_x0000_t136"'
        '          style="position:absolute;margin-left:0;margin-top:0;'
        '                 width:467.85pt;height:228.8pt;'
        '                 rotation:315;'
        '                 z-index:-251654144;mso-position-horizontal:center;'
        '                 mso-position-horizontal-relative:margin;'
        '                 mso-position-vertical:center;'
        '                 mso-position-vertical-relative:margin"'
        '          fillcolor="#d0d0d0" stroked="f">'
        '        <v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt;font-weight:bold"'
        '            string="DRAFT" trim="t" fitshape="t"/>'
        '      </v:shape>'
        '    </w:pict>'
        '  </w:r>'
        '</w:p>'
    )
    wmark_el = etree.fromstring(wmark_xml)
    hdr_xml.append(wmark_el)

    # Word requires at least one paragraph after the watermark element
    empty = OxmlElement('w:p')
    hdr_xml.append(empty)


# ── Cover page section helpers ────────────────────────────────────────────────

def add_cover_page_border(section, offset_pt: int = 18, sz: int = 18, color: str = "999999"):
    """
    Add a w:pgBorders element to the section's sectPr.

    offset_pt : distance from paper edge in points (18pt ≈ 0.25")
    sz        : border line thickness in 1/8-point units (18 = ~2.25pt)
    color     : 6-char hex string without '#'

    Placement: pgBorders is inserted after w:pgMar in sectPr to satisfy
    the OOXML schema ordering requirement.
    """
    sectPr = section._sectPr
    for existing in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(existing)

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')
    for side in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    str(sz))
        el.set(qn('w:space'), str(offset_pt))
        el.set(qn('w:color'), color)
        pgBorders.append(el)

    pgMar = sectPr.find(qn('w:pgMar'))
    if pgMar is not None:
        pgMar.addnext(pgBorders)
    else:
        sectPr.append(pgBorders)


def build_cover_footer(section, meta: dict):
    """
    Build the cover page footer with the org copyright / contact block.

    The footer is always rendered at the bottom margin by Word, so this
    block is reliably positioned regardless of cover page content length
    and cannot overlap the body watermark.

    Org identity is read from meta['org'] (a nested dict) with fallbacks
    to the module-level constants so existing documents without an org
    block continue to work unchanged.

    Supported meta['org'] keys:
        name   — organisation name     (default: ORG_NAME)
        dept   — department / unit     (default: ORG_DEPT)
        addr1  — street / PO box       (default: ORG_ADDR1)
        addr2  — city / state / zip    (default: ORG_ADDR2)
        url    — website URL           (default: ORG_URL)

    Example front matter:
        org:
          name:  "Acme Corp"
          dept:  "Enterprise Architecture"
          addr1: "123 Main Street"
          addr2: "Anytown, USA"
          url:   "www.example.com"
    """
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear any inherited paragraphs
    for para in list(footer.paragraphs):
        p = para._p
        try:
            p.getparent().remove(p)
        except Exception:
            pass

    # Read org identity — meta['org'] overrides constants
    org   = meta.get('org') or {}
    name  = org.get('name',  ORG_NAME)
    dept  = org.get('dept',  ORG_DEPT)
    addr1 = org.get('addr1', ORG_ADDR1)
    addr2 = org.get('addr2', ORG_ADDR2)
    url   = org.get('url',   ORG_URL)

    def _footer_line(text, bold=False, is_url=False):
        p = footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=0, after=0)
        run = p.add_run(text)
        run.bold      = bold
        run.font.name = FONT_BODY
        run.font.size = Pt(9)
        set_run_color(run, BLUE_LINK if is_url else DARK_NAVY)
        if is_url:
            run.underline = True

    _footer_line(f'\u00A9{name}', bold=True)   # © OrgName
    _footer_line(dept)
    _footer_line(addr1)
    _footer_line(addr2)
    _footer_line(url, is_url=True)


def remove_page_border(section):
    """Remove all w:pgBorders from a section so it does not inherit the cover border."""
    sectPr = section._sectPr
    for el in sectPr.findall(qn('w:pgBorders')):
        sectPr.remove(el)


def set_toc_page_numbering_start(section, start: int = 1):
    """
    Insert w:pgNumType to restart page numbering at `start` in this section.
    Used to reset numbering to 1 at the TOC section so body pages count from 1.
    Inserted before w:cols in sectPr to satisfy schema ordering.
    """
    sectPr = section._sectPr
    for el in sectPr.findall(qn('w:pgNumType')):
        sectPr.remove(el)
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), str(start))
    cols = sectPr.find(qn('w:cols'))
    if cols is not None:
        cols.addprevious(pgNumType)
    else:
        sectPr.append(pgNumType)


# ── Cover page builder ────────────────────────────────────────────────────────

def build_cover_page(doc, meta: dict, logo_path: str | None):
    """
    Build and populate the cover page (document section 1).

    Layout (top to bottom):
      - Logo (image) or org name fallback text, left-aligned
      - Thin rule under logo
      - Vertical spacer (~10 empty paragraphs)
      - Document title in UPPER CASE, large bold
      - Thin rule under title
      - Department / org unit, italic
      [section footer — always at bottom margin]
      - Copyright block: org name, dept, address, URL — centered, small

    The copyright block is rendered in the Word section footer (not the
    document body) so it is always pinned to the bottom margin regardless
    of title length or department line count, and cannot overlap the
    watermark or shift with content changes.

    If status == 'Draft' (case-insensitive), a diagonal DRAFT watermark
    is added to the section header.

    Spacing spacers: the 10 empty-paragraph count is tuned to the reference
    document. Adjust here if the vertical balance of the title block needs
    tuning.
    """
    section = doc.sections[0]
    section.page_width  = Emu(PAGE_W_TWIPS * TWIPS_TO_EMU)
    section.page_height = Emu(PAGE_H_TWIPS * TWIPS_TO_EMU)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

    add_cover_page_border(section, offset_pt=18, sz=18, color="999999")

    if meta.get('status', '').lower() == 'draft':
        add_draft_watermark(section)
    else:
        hdr = section.header
        hdr.is_linked_to_previous = False

    # Copyright block lives in the section footer so it is always pinned to
    # the bottom margin, independent of body content length.
    build_cover_footer(section, meta)

    # ── Local helper: add a styled paragraph to the document body ─────────────
    def doc_para(text='', bold=False, italic=False, size=11,
                 color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT,
                 before=0, after=0, font=FONT_BODY):
        p = doc.add_paragraph()
        p.alignment = align
        para_spacing(p, before=before, after=after)
        if text:
            run = p.add_run(text)
            run.bold      = bold
            run.italic    = italic
            run.font.name = font
            run.font.size = Pt(size)
            set_run_color(run, color)
        return p

    # ── Logo block ────────────────────────────────────────────────────────────
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(logo_para, before=120, after=0)

    if logo_path and os.path.isfile(logo_path):
        run = logo_para.add_run()
        run.add_picture(logo_path, width=Inches(2.2))
    else:
        # Fallback: org name as styled text when no logo file is provided
        run = logo_para.add_run(ORG_NAME)
        run.bold = True
        run.font.name = FONT_TITLE
        run.font.size = Pt(14)
        set_run_color(run, BLUE_DARK)

    # Rule under logo
    rule_para = doc.add_paragraph()
    para_spacing(rule_para, before=60, after=0)
    set_paragraph_border_bottom(rule_para, color="1F4E79", sz=6)

    # ── Vertical spacer: push title toward vertical center ────────────────────
    for _ in range(10):
        doc_para(before=0, after=0)

    # ── Document title ────────────────────────────────────────────────────────
    # Title supports embedded newlines (literal \n in YAML or actual newlines)
    title_text  = meta.get('title', 'Untitled Document')
    title_lines = title_text.split('\\n') if '\\n' in title_text else title_text.split('\n')

    for line in title_lines:
        line = line.strip()
        if not line:
            continue
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_spacing(tp, before=0, after=60)
        run = tp.add_run(line.upper())
        run.bold      = True
        run.font.name = FONT_TITLE
        run.font.size = Pt(20)
        set_run_color(run, BLACK)

    # Rule under title
    rule_para2 = doc.add_paragraph()
    para_spacing(rule_para2, before=100, after=100)
    set_paragraph_border_bottom(rule_para2, color="000000", sz=4)

    # ── Department line ───────────────────────────────────────────────────────
    dept = meta.get('department', '')
    if dept:
        doc_para(dept, italic=True, size=12, color=BLACK, before=80, after=80)

    # (Copyright block is in the section footer — see build_cover_footer above)
