"""
revision_history.py — Revision history table and document metadata block.

Renders on its own page (section 3). Contains:
  - A styled "Revision History" heading
  - A table with Version / Date / Author / Description columns
  - A metadata block below the table: Owner, Audience, Related Docs

If no revision_history list is present in front matter, a single row is
auto-generated from the top-level version/date/author fields.
"""

from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor

from .constants import BLUE_DARK, BLACK, FONT_BODY, FONT_TITLE
from .xml_helpers import (
    set_run_color, para_spacing,
    set_cell_bg, set_cell_borders, set_table_border,
)


def build_revision_table(doc, meta: dict):
    """
    Add a revision history table to the document.

    Args:
        doc:  python-docx Document
        meta: parsed YAML front matter dictionary

    Table columns: Version | Date | Author | Description
    Column widths sum to 8.0 inches (matches 1-inch margins on 8.5" page).
    Header row has a dark navy background (#1F4E79) with white text.
    Body rows alternate white / light gray (#F2F2F2).
    """
    # ── Section heading ───────────────────────────────────────────────────────
    heading = doc.add_paragraph()
    para_spacing(heading, before=0, after=240)
    run = heading.add_run("Revision History")
    run.bold      = True
    run.font.name = FONT_TITLE
    run.font.size = Pt(14)
    set_run_color(run, BLUE_DARK)

    cols   = ["Version", "Date", "Author", "Description"]
    widths = [0.8, 1.2, 1.5, 4.5]   # inches; sum = 8.0

    table = doc.add_table(rows=1, cols=4)
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(table, color="CCCCCC", sz=4)

    # ── Header row ────────────────────────────────────────────────────────────
    hdr = table.rows[0]
    hdr.height = Cm(0.7)
    for i, (col, w) in enumerate(zip(cols, widths)):
        cell = hdr.cells[i]
        cell.width = Inches(w)
        set_cell_bg(cell, "1F4E79")
        set_cell_borders(cell, color="CCCCCC", sz=4)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_spacing(p, before=60, after=60)
        run = p.add_run(col)
        run.bold      = True
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        set_run_color(run, RGBColor(0xFF, 0xFF, 0xFF))

    # ── Revision rows ─────────────────────────────────────────────────────────
    revisions = meta.get('revision_history', [])
    if not revisions:
        # Auto-generate from top-level metadata fields when no explicit list
        revisions = [{
            'version':     meta.get('version', '1.0'),
            'date':        str(meta.get('date', '')),
            'author':      meta.get('author', ''),
            'description': 'Initial draft' if meta.get('status', '').lower() == 'draft'
                           else 'Initial release',
        }]

    for idx, rev in enumerate(revisions):
        row  = table.add_row()
        fill = "F2F2F2" if idx % 2 == 1 else "FFFFFF"
        vals = [
            str(rev.get('version', '')),
            str(rev.get('date',    '')),
            str(rev.get('author',  '')),
            str(rev.get('description', '')),
        ]
        for i, (val, w) in enumerate(zip(vals, widths)):
            cell = row.cells[i]
            cell.width = Inches(w)
            set_cell_bg(cell, fill)
            set_cell_borders(cell, color="CCCCCC", sz=4)
            p = cell.paragraphs[0]
            para_spacing(p, before=60, after=60)
            run = p.add_run(val)
            run.font.name = FONT_BODY
            run.font.size = Pt(10)
            set_run_color(run, BLACK)

    # ── Metadata block below table ────────────────────────────────────────────
    audience     = meta.get('audience', [])
    related_docs = meta.get('related_docs', [])

    meta_fields = [
        ("Document Owner", meta.get('owner', '')),
        ("Audience",       ', '.join(audience)     if isinstance(audience, list)     else str(audience)),
        ("Related Docs",   ', '.join(related_docs) if isinstance(related_docs, list) else str(related_docs)),
    ]

    doc.add_paragraph()  # spacer between table and metadata block
    for label, value in meta_fields:
        if not value:
            continue
        p = doc.add_paragraph()
        para_spacing(p, before=40, after=40)
        label_run = p.add_run(f"{label}: ")
        label_run.bold      = True
        label_run.font.name = FONT_BODY
        label_run.font.size = Pt(10)
        set_run_color(label_run, BLUE_DARK)
        val_run = p.add_run(value)
        val_run.font.name = FONT_BODY
        val_run.font.size = Pt(10)
        set_run_color(val_run, BLACK)
