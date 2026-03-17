"""
toc.py — Table of Contents page rendering.

FRAGILITY WARNING
-----------------
The TOC is built in two layers:

1. Native Word TOC field (w:fldChar / w:instrText)
   This is the "real" TOC that Word populates automatically when the document
   is opened. It uses heading styles (Heading1–Heading3) to discover entries
   and fills in actual page numbers. The field instruction is:
       TOC \\o "1-3" \\h \\z \\u
   where:  \\o "1-3" = outline levels 1-3
           \\h       = make entries hyperlinks
           \\z       = hide tab leader and page numbers in web layout
           \\u       = use applied paragraph outline level

2. Static fallback entries
   These are pre-rendered paragraphs with dot-leader tab stops. They are
   visible before Word updates the field (e.g., when the file is opened for
   the first time in some viewers). They use placeholder page numbers and
   will be stale after the first Word update, which is expected.

Both layers must remain in the document. The native field is the source of
truth; the static entries are cosmetic scaffolding.

Tab stop note: dot-leader tab stops (w:tab val="right" leader="dot") must
be set via raw XML because python-docx does not expose leader styles through
its public API. The position 8640 twips = 6 inches from the left margin.
"""

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .constants import TEAL, BLACK, FONT_BODY, FONT_TITLE
from .xml_helpers import set_run_color, para_spacing


def build_toc_page(doc, headings: list[tuple]):
    """
    Build the Table of Contents page.

    Args:
        doc:      python-docx Document
        headings: list of (level: int, text: str, page: int | None)
                  Typically pre-scanned from the markdown source.
                  page values are placeholders; Word updates them on open.

    The function adds paragraphs directly to the document body. It does not
    add a section break — the caller is responsible for section management.
    """
    # ── TOC heading ───────────────────────────────────────────────────────────
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para_spacing(toc_heading, before=0, after=240)
    run = toc_heading.add_run("Table of Contents")
    run.bold = True
    run.font.name = FONT_TITLE
    run.font.size = Pt(14)
    set_run_color(run, TEAL)

    # ── Native Word TOC field ─────────────────────────────────────────────────
    # Word populates this field with actual headings and page numbers on open.
    # Do not remove or reorder the begin / instrText / end runs.
    toc_para = doc.add_paragraph()
    toc_para._p.clear()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run_begin = OxmlElement('w:r')
    run_begin.append(fldChar_begin)
    run_instr = OxmlElement('w:r')
    run_instr.append(instrText)
    run_end   = OxmlElement('w:r')
    run_end.append(fldChar_end)

    toc_para._p.append(run_begin)
    toc_para._p.append(run_instr)
    toc_para._p.append(run_end)

    # ── Static fallback entries ───────────────────────────────────────────────
    # Rendered as dot-leader paragraphs. Page numbers are placeholders.
    # These are shown before Word updates the TOC field. They do not need
    # to stay in sync after the first open — that is Word's job.
    indent_per_level = {1: 0.0, 2: 0.25, 3: 0.50}

    for level, text, page in headings:
        toc_entry = doc.add_paragraph()
        toc_entry.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_spacing(toc_entry, before=0, after=60)

        indent = indent_per_level.get(level, 0.0)
        if indent:
            toc_entry.paragraph_format.left_indent = Inches(indent)

        # Dot-leader right-aligned tab stop at 8640 twips (6" from left margin).
        # This must be set via raw XML — python-docx has no public API for leaders.
        pPr  = toc_entry._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab  = OxmlElement('w:tab')
        tab.set(qn('w:val'),    'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'),    '8640')
        tabs.append(tab)
        # w:tabs must appear before w:spacing and w:ind in pPr
        pPr.insert(0, tabs)

        # Entry text run
        run = toc_entry.add_run(text)
        run.font.name = FONT_BODY
        run.font.size = Pt(10)
        run.bold      = (level == 1)
        set_run_color(run, BLACK)

        # Tab character + page number
        tab_run = toc_entry.add_run()
        tab_run.font.name = FONT_BODY
        tab_run.font.size = Pt(10)
        br = OxmlElement('w:tab')
        tab_run._r.append(br)

        page_run = toc_entry.add_run(str(page) if page else ' ')
        page_run.font.name = FONT_BODY
        page_run.font.size = Pt(10)
        set_run_color(page_run, BLACK)
