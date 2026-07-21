"""
diagrams.py — Diagram extraction and rendering for docx_builder.

Pipeline role
-------------
Inline diagram fences — ```mermaid blocks plus any fence whose info string is
a Kroki-supported diagram language (```plantuml, ```graphviz, ```d2, ...) —
are extracted from the markdown source before the mistune pass and replaced
with __MERMAID_N__ placeholders. (The token name predates multi-language
support and is kept for compatibility; it covers all diagram languages.)
builder.py calls render_diagram() for each placeholder while iterating body
segments in source order, so diagrams land in the correct position relative
to surrounding paragraphs.

Fences whose info string is NOT a supported diagram language (```bash,
```yaml, ```python, ...) are left completely untouched and continue to
render as ordinary code blocks.

Rendering backends
------------------
Mermaid fences try backends in this order:

  1. mmdc       — Mermaid CLI (Node.js). Best output quality and the only
                  backend that supports custom CSS themes. Used automatically
                  when mmdc is found in PATH (GitHub Actions, devspace).
                  Set DOCX_BUILDER_MERMAID_BACKEND=mmdc to require it.

  2. Kroki      — HTTP API (https://kroki.io by default, or a self-hosted
                  instance via DOCX_BUILDER_KROKI_URL). Diagram source is
                  POSTed to {kroki_url}/{diagram_type}/png. No local
                  dependencies; requires network access at build time.

  3. mermaid.ink — Public HTTP API at https://mermaid.ink (secondary API).
                   Tried if Kroki fails. Uses plain base64url encoding.

  4. Placeholder — If all backends fail (offline, syntax error, no Node.js),
                   a styled warning block containing the raw diagram source is
                   inserted instead of crashing the build.

All other diagram languages render via Kroki only (backend 2), falling back
to the styled placeholder (backend 4) on any failure.

Environment variable overrides:
    DOCX_BUILDER_MERMAID_BACKEND=auto        (default — mmdc → kroki → ink)
    DOCX_BUILDER_MERMAID_BACKEND=mmdc        (require mmdc only)
    DOCX_BUILDER_MERMAID_BACKEND=kroki       (skip mmdc, Kroki only)
    DOCX_BUILDER_MERMAID_BACKEND=mermaid_ink (skip mmdc and kroki, ink only)

    DOCX_BUILDER_KROKI_URL=https://kroki.io  (default — public Kroki API)
    DOCX_BUILDER_KROKI_URL=http://127.0.0.1:8000
                                             (self-hosted Kroki instance;
                                             applies to ALL diagram languages,
                                             including the mermaid Kroki step)

Custom CSS (mmdc only)
----------------------
Provide a CSS file via the DOCX_BUILDER_MERMAID_CSS environment variable, or
pass an explicit css_path argument from the caller. When mmdc is the active
backend, this file is passed via --cssFile. Ignored by mermaid.ink.

Supported image formats for ![]() references
---------------------------------------------
python-docx supports: PNG, JPG/JPEG, GIF, TIFF, BMP, EMF, WMF.
SVG is NOT supported natively — export to PNG before referencing.

Caption syntax for inline diagram fences
-----------------------------------------
    ```mermaid caption="Figure 1 — DevOps Lifecycle"
    flowchart LR
        A --> B
    ```

The caption attribute is optional and works on any diagram fence
(```plantuml caption="..." etc.). Alt text in ![]() image references is
used as the caption for path-referenced images.
"""

import os
import re
import base64
import shutil
import urllib.request
import urllib.error
import urllib.parse
import hashlib

from docx.shared import Inches, Pt
from docx.image.image import Image as DocxImage
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .constants import FONT_BODY, DARK_NAVY, GRAY_TEXT, BLACK

# ── Module-level defaults ─────────────────────────────────────────────────────

# Default rendered width in the document body (inches).
# Fits within 1" margins on 8.5" US Letter with some breathing room.
DIAGRAM_DEFAULT_WIDTH_IN: float = 5.5
DIAGRAM_MAX_HEIGHT_IN: float = 4.75

# mermaid.ink theme applied to all diagrams when using the API backend.
# Options: 'default' | 'neutral' | 'forest' | 'dark'
# 'neutral' uses greyscale tones appropriate for formal documentation.
# 'default' uses blues that are closer to the default document colour scheme.
# Can be overridden via front matter: diagrams.mermaid_theme
MERMAID_INK_DEFAULT_THEME: str = 'neutral'

# Render diagrams at a higher pixel density than the final on-page
# size so dense node/link layouts stay readable when embedded into DOCX.
DIAGRAM_RENDER_SCALE_DEFAULT: float = 2.0
DIAGRAM_RENDER_SCALE_MIN: float = 1.0
DIAGRAM_RENDER_SCALE_MAX: float = 4.0

# Default Kroki endpoint. Override with DOCX_BUILDER_KROKI_URL to point at a
# self-hosted Kroki instance (e.g. http://127.0.0.1:8000).
DEFAULT_KROKI_URL: str = 'https://kroki.io'

# Fence info string → Kroki diagram type.
# Keys are the languages recognised as diagram fences; values are the path
# segment used in the Kroki render URL ({kroki_url}/{type}/png). Most map
# 1:1; aliases like ```dot → graphviz are the exception.
#
# KEEP IN STEP with SUPPORTED_DIAGRAM_TYPES in scripts/docx_manifest.py —
# scripts/ is not an importable package, so the list is duplicated there.
# Any fence whose info string is NOT in this map is an ordinary code block
# (```bash, ```yaml, ```python, ...) and must never be treated as a diagram.
KROKI_LANGUAGE_TYPES: dict[str, str] = {
    'mermaid': 'mermaid',
    'plantuml': 'plantuml',
    'c4plantuml': 'c4plantuml',
    'graphviz': 'graphviz',
    'dot': 'graphviz',
    'd2': 'd2',
    'pikchr': 'pikchr',
    'erd': 'erd',
    'svgbob': 'svgbob',
    'nomnoml': 'nomnoml',
    'structurizr': 'structurizr',
    'ditaa': 'ditaa',
    'seqdiag': 'seqdiag',
    'blockdiag': 'blockdiag',
    'nwdiag': 'nwdiag',
    'packetdiag': 'packetdiag',
    'rackdiag': 'rackdiag',
    'umlet': 'umlet',
    'vega': 'vega',
    'vegalite': 'vegalite',
    'wavedrom': 'wavedrom',
    'wireviz': 'wireviz',
    'dbml': 'dbml',
    'bpmn': 'bpmn',
    'excalidraw': 'excalidraw',
    'bytefield': 'bytefield',
    'goat': 'goat',
    'symbolator': 'symbolator',
}


# ── Extraction ────────────────────────────────────────────────────────────────

def extract_diagram_fences(md_text: str) -> tuple[str, list[dict]]:
    """
    Extract diagram code fences from markdown and replace with placeholders.

    A fence is a diagram fence when its info string (the word after the
    opening backticks) is a key in KROKI_LANGUAGE_TYPES — ```mermaid,
    ```plantuml, ```graphviz, ```d2, etc. Any other fence (```bash,
    ```yaml, ```python, plain ```) is NOT a diagram: it is returned to the
    markdown stream byte-for-byte unchanged and renders as an ordinary
    code block.

    Returns (processed_text, diagram_data_list) where:
      - processed_text has each diagram fence replaced by a __MERMAID_N__
        token (name kept for compatibility; covers all diagram languages)
      - diagram_data_list[N] is a dict:
            {'source': str, 'caption': str | None,
             'language': str, 'kroki_type': str}

    The caller (builder.py) splits the processed text on __MERMAID_N__ and
    __TABLE_N__ tokens and calls render_diagram() for each diagram token,
    keeping all elements in source order.

    Caption is parsed from an optional caption="..." attribute on the opening
    fence line:
        ```mermaid caption="Figure 1 — DevOps Lifecycle"
    """
    # Matches ```<language><optional attrs>\n<body>\n``` — non-greedy body.
    # The language must start at the fence with no gap; whether it is a
    # diagram is decided in _extract so non-diagram fences pass through.
    fence_re = re.compile(
        r'^```([A-Za-z][A-Za-z0-9_-]*)([^\n]*)\n(.*?)^```',
        re.MULTILINE | re.DOTALL,
    )
    diagrams: list[dict] = []

    def _extract(m: re.Match) -> str:
        language = m.group(1).strip().lower()
        if language not in KROKI_LANGUAGE_TYPES:
            # Ordinary code block (```bash, ```yaml, ...) — leave untouched.
            return m.group(0)

        attrs_str = m.group(2).strip()
        source    = m.group(3).strip()

        cap_match = re.search(r'caption=["\']([^"\']+)["\']', attrs_str)
        caption   = cap_match.group(1) if cap_match else None

        idx = len(diagrams)
        diagrams.append({
            'source': source,
            'caption': caption,
            'language': language,
            'kroki_type': KROKI_LANGUAGE_TYPES[language],
        })
        return f'\n__MERMAID_{idx}__\n'

    processed = fence_re.sub(_extract, md_text)
    return processed, diagrams


# Backwards-compatible alias — the function handled only ```mermaid fences
# before multi-language Kroki support landed. Existing callers keep working.
extract_mermaid_fences = extract_diagram_fences


# ── Rendering ─────────────────────────────────────────────────────────────────

def render_diagram(
    doc,
    diagram_data: dict,
    tmp_dir: str,
    mermaid_theme: str = MERMAID_INK_DEFAULT_THEME,
    css_path: str | None = None,
) -> None:
    """
    Render a single diagram into the document at the current position.

    Mermaid diagrams try mmdc first (if in PATH or forced via env var), then
    Kroki, then mermaid.ink. Every other diagram language renders via Kroki
    only. Both paths fall back to a styled placeholder so the build never
    hard-fails.

    Args:
        doc:           The python-docx Document object.
        diagram_data:  Dict with keys 'source' (str) and 'caption' (str|None).
                       Optional 'language' and 'kroki_type' keys (added by
                       extract_diagram_fences) select the render path;
                       both default to 'mermaid' for older callers.
        tmp_dir:       Temporary directory for intermediate PNG files.
        mermaid_theme: mermaid.ink theme string (ignored by mmdc — use css_path).
        css_path:      Optional path to a .css file passed to mmdc --cssFile.
    """
    source     = diagram_data['source']
    caption    = diagram_data.get('caption')
    language   = diagram_data.get('language', 'mermaid')
    kroki_type = diagram_data.get('kroki_type', KROKI_LANGUAGE_TYPES.get(language, language))
    render_scale = _get_mermaid_render_scale()

    digest = hashlib.sha256(source.encode("utf-8")).hexdigest()
    png_path = os.path.join(tmp_dir, f'{language}_{digest}.png')
    rendered = False

    if language == 'mermaid':
        backend = os.environ.get('DOCX_BUILDER_MERMAID_BACKEND', 'auto').lower()

        # 1. mmdc — best quality, supports custom CSS; used when in PATH
        if backend in ('auto', 'mmdc') and shutil.which('mmdc'):
            _css = css_path or os.environ.get('DOCX_BUILDER_MERMAID_CSS')
            rendered = _render_via_mmdc(
                source,
                png_path,
                css_path=_css,
                scale=render_scale,
            )

        # 2. Kroki — primary API backend, generally more reliable
        if not rendered and backend in ('auto', 'kroki'):
            rendered = _render_via_kroki(source, png_path, scale=render_scale)

        # 3. mermaid.ink — secondary public API fallback
        if not rendered and backend in ('auto', 'mermaid_ink'):
            rendered = _render_via_mermaid_ink(
                source,
                png_path,
                theme=mermaid_theme,
                scale=render_scale,
            )
    else:
        # Non-mermaid languages have exactly one backend: Kroki.
        rendered = _render_via_kroki(
            source,
            png_path,
            diagram_type=kroki_type,
            scale=render_scale,
        )

    if rendered and os.path.isfile(png_path):
        _embed_image(doc, png_path, caption)
    else:
        _embed_placeholder(doc, source, caption, language=language)


# ── Private: rendering backends ───────────────────────────────────────────────

def _get_kroki_url() -> str:
    """Return the Kroki base URL from the environment (default: kroki.io)."""
    raw = os.environ.get('DOCX_BUILDER_KROKI_URL', '').strip()
    return (raw or DEFAULT_KROKI_URL).rstrip('/')


def _render_via_kroki(
    source: str,
    output_path: str,
    diagram_type: str = 'mermaid',
    scale: float = DIAGRAM_RENDER_SCALE_DEFAULT,
) -> bool:
    """
    Fetch a rendered PNG from a Kroki server.

    The diagram source is POSTed as text/plain to
    {kroki_url}/{diagram_type}/png — the same form scripts/docx_manifest.py
    uses, and it avoids URL-length limits that the GET encoding hits on
    large diagrams. The endpoint defaults to the public https://kroki.io
    API and can point at a self-hosted instance via DOCX_BUILDER_KROKI_URL.
    Supports Mermaid and every other Kroki diagram type. No local
    dependencies.

    Returns True on success, False on any network or API error.
    """
    try:
        params = urllib.parse.urlencode({'scale': _format_scale_query_value(scale)})
        url = f'{_get_kroki_url()}/{diagram_type}/png?{params}'
        req = urllib.request.Request(
            url,
            data=source.encode('utf-8'),
            headers={
                'User-Agent': 'docx-builder/1.0',
                'Content-Type': 'text/plain; charset=utf-8',
            },
            method='POST',
        )
        with urllib.request.urlopen(req, timeout=20) as resp:
            content_type = resp.headers.get('Content-Type', '')
            if 'image' not in content_type:
                print(f'  [diagrams] Kroki returned non-image '
                      f'content-type: {content_type}')
                return False
            with open(output_path, 'wb') as f:
                f.write(resp.read())
        return True
    except (urllib.error.URLError, OSError) as exc:
        print(f'  [diagrams] Kroki render failed ({diagram_type}): {exc}')
        return False


def _render_via_mermaid_ink(
    source: str,
    output_path: str,
    theme: str = MERMAID_INK_DEFAULT_THEME,
    scale: float = DIAGRAM_RENDER_SCALE_DEFAULT,
) -> bool:
    """
    Fetch a rendered PNG from the mermaid.ink public API.

    Encodes the Mermaid source as URL-safe base64 and retrieves a PNG.
    Requires internet access at build time. No local runtime dependencies.

    Returns True on success, False on any network or API error.
    """
    try:
        encoded = base64.urlsafe_b64encode(source.encode('utf-8')).decode('ascii')
        # bgColor=!white keeps the diagram background white regardless of theme
        params = urllib.parse.urlencode({
            'bgColor': '!white',
            'theme': theme,
            'scale': _format_scale_query_value(scale),
        })
        url = f'https://mermaid.ink/img/{encoded}?{params}'
        req = urllib.request.Request(
            url,
            headers={'User-Agent': 'docx-builder/1.0'},
        )
        with urllib.request.urlopen(req, timeout=20) as resp:
            content_type = resp.headers.get('Content-Type', '')
            if 'image' not in content_type:
                # mermaid.ink returns an HTML error page for invalid syntax
                print(f'  [diagrams] mermaid.ink returned non-image '
                      f'content-type: {content_type}')
                return False
            with open(output_path, 'wb') as f:
                f.write(resp.read())
        return True
    except (urllib.error.URLError, OSError) as exc:
        print(f'  [diagrams] mermaid.ink render failed: {exc}')
        return False


def _render_via_mmdc(
    source: str,
    output_path: str,
    css_path: str | None = None,
    scale: float = DIAGRAM_RENDER_SCALE_DEFAULT,
) -> bool:
    """
    Render a Mermaid diagram via the mmdc CLI (requires Node.js).

    mmdc must be available in PATH. This backend is preferred when available
    because it produces higher-quality output and supports custom CSS themes.

    Intended for use in:
      - GitHub Actions (install Node + @mermaid-js/mermaid-cli as a pre-step)
      - OpenShift devspace containers with Node.js included
      - Local development when Node.js is installed

    css_path: Optional path to a CSS file passed to mmdc --cssFile.
              See diagrams/mermaid-theme.css for a starter theme file.

    Returns True on success, False if mmdc is missing or exits with an error.
    """
    import subprocess
    import tempfile as _tf

    mmd_path = None
    try:
        with _tf.NamedTemporaryFile(
            mode='w', suffix='.mmd', delete=False, encoding='utf-8'
        ) as tmp:
            tmp.write(source)
            mmd_path = tmp.name

        cmd = [
            'mmdc',
            '-i',
            mmd_path,
            '-o',
            output_path,
            '-b',
            'white',
            '-s',
            _format_scale_query_value(scale),
        ]
        if css_path and os.path.isfile(css_path):
            cmd += ['--cssFile', css_path]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

        if result.returncode == 0 and os.path.isfile(output_path):
            return True
        print(f'  [diagrams] mmdc exited {result.returncode}: '
              f'{result.stderr.strip()[:200]}')
        return False
    except (subprocess.TimeoutExpired, FileNotFoundError, OSError) as exc:
        print(f'  [diagrams] mmdc render failed: {exc}')
        return False
    finally:
        if mmd_path and os.path.exists(mmd_path):
            try:
                os.unlink(mmd_path)
            except OSError:
                # Best-effort cleanup; ignore errors during unlink.
                pass


# ── Private: document embedding ───────────────────────────────────────────────

def _embed_image(doc, img_path: str, caption: str | None) -> None:
    """Embed a PNG centred in the document body with an optional italic caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    from .xml_helpers import para_spacing
    para_spacing(p, before=120, after=40)
    width, height = fit_image_dimensions(img_path)
    if height is None:
        p.add_run().add_picture(img_path, width=width)
    else:
        p.add_run().add_picture(img_path, width=width, height=height)

    if caption:
        _add_caption(doc, caption)


def _embed_placeholder(doc, source: str, caption: str | None,
                       language: str = 'mermaid') -> None:
    """
    Emit a styled warning block when diagram rendering fails.

    Preserves the raw diagram source in a code-style paragraph so the
    document author can see the intent. Does not raise an exception.
    The hint text names the language and its remedy: mermaid suggests the
    CLI install; every other language points at the Kroki endpoint.
    """
    from .xml_helpers import para_spacing, set_run_color

    if language == 'mermaid':
        hint = ('\u26a0 Diagram could not be rendered '
                '(install @mermaid-js/mermaid-cli or check internet access)')
    else:
        hint = (f'\u26a0 {language} diagram could not be rendered '
                f'(check Kroki endpoint: {_get_kroki_url()})')

    # Warning banner
    warn = doc.add_paragraph()
    para_spacing(warn, before=120, after=20)
    run = warn.add_run(hint)
    run.bold      = True
    run.italic    = True
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    set_run_color(run, DARK_NAVY)

    # Diagram source in code style with grey background
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    para_spacing(p, before=0, after=120)
    code_run = p.add_run(source)
    code_run.font.name = 'Courier New'
    code_run.font.size = Pt(8)
    set_run_color(code_run, BLACK)

    if caption:
        _add_caption(doc, f'[{caption}]')


def _add_caption(doc, text: str) -> None:
    """Add a small centred italic caption paragraph below a diagram."""
    from .xml_helpers import para_spacing, set_run_color

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_spacing(cap, before=0, after=120)
    run           = cap.add_run(text)
    run.italic    = True
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    set_run_color(run, GRAY_TEXT)


def fit_image_dimensions(
    img_path: str,
    *,
    max_width_in: float = DIAGRAM_DEFAULT_WIDTH_IN,
    max_height_in: float = DIAGRAM_MAX_HEIGHT_IN,
):
    """Return width/height bounded to a page-friendly box while preserving aspect ratio."""
    image = DocxImage.from_file(img_path)
    width_in = image.width / 914400
    height_in = image.height / 914400
    if width_in <= 0 or height_in <= 0:
        return Inches(max_width_in), None

    scale = min(max_width_in / width_in, max_height_in / height_in, 1.0)
    return Inches(width_in * scale), Inches(height_in * scale)


def _get_mermaid_render_scale() -> float:
    """Return a bounded diagram render scale from the environment."""
    raw = ''
    for name in (
        'DOCX_BUILDER_DIAGRAM_RENDER_SCALE',
        'DOCX_BUILDER_MERMAID_RENDER_SCALE',
    ):
        value = os.environ.get(name)
        if value is not None and value.strip():
            raw = value.strip()
            break
    if not raw:
        raw = str(DIAGRAM_RENDER_SCALE_DEFAULT)
    try:
        scale = float(raw)
    except ValueError:
        return DIAGRAM_RENDER_SCALE_DEFAULT

    if scale < DIAGRAM_RENDER_SCALE_MIN:
        return DIAGRAM_RENDER_SCALE_MIN
    if scale > DIAGRAM_RENDER_SCALE_MAX:
        return DIAGRAM_RENDER_SCALE_MAX
    return scale


def _format_scale_query_value(scale: float) -> str:
    """Serialize scale values cleanly for CLI flags and HTTP query strings."""
    if float(scale).is_integer():
        return str(int(scale))
    return str(scale)
