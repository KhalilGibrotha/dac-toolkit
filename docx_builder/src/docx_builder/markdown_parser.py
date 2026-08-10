"""
markdown_parser.py — Markdown to python-docx paragraph conversion.

Pipeline:
  1. extract_mermaid_fences() (diagrams.py) — extract ```mermaid blocks,
                          replace with __MERMAID_N__ placeholders.
  2. extract_md_tables()  — extract GFM-style pipe tables, render them as
                          docx tables, and replace them with placeholders
                          in the markdown text.
  3. mistune renders remaining markdown to HTML.
  4. HtmlToDocx (HTMLParser subclass) walks the HTML and emits paragraphs.

Why the multi-stage approach:
  Mermaid fences and pipe tables must be extracted before mistune sees them
  so they can be rendered with full control over styling and placement.
  Diagrams and tables are rendered inline during the segment loop in
  builder.py, so they appear in the correct document position.

Supported markdown elements:
  Block:  h1–h6, p, ul (unordered), ol (ordered), blockquote, pre/code, hr
  Inline: strong, em, code, a (live hyperlink for absolute http/https/mailto/
          ftp targets; relative paths and bare #fragments render as styled
          text, since neither has an absolute form to point a reader at)
  Images: ![alt text](relative/path/to/image.png) — resolved relative to
          the source .md file's directory (PNG, JPG, GIF, TIFF, BMP)
  Tables: GFM pipe tables with header / alignment divider / body rows

Elements not supported (rendered as plain text or ignored):
  - h4–h6 styled as h3 (clamped — template defines three heading levels)
  - Nested blockquotes beyond one level
  - HTML inside markdown (stripped)
  - SVG images (export to PNG before referencing)
  - Cover logo (handled separately in cover_page.py)
"""

import os
import re
from html.parser import HTMLParser
from urllib.parse import quote

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import RGBColor
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .constants import (BLUE_LINK, BLACK, FONT_BODY, FONT_MONO, GRAY_TEXT,
                        SIZE_BODY, SIZE_CAPTION, SIZE_CODE,
                        SIZE_CODE_IN_CELL, SIZE_TABLE_CELL, SIZE_TABLE_HEADER,
                        TABLE_COL_CHAR_CAP, TABLE_MIN_COL_IN,
                        TABLE_TOTAL_WIDTH_IN)
from .diagrams import fit_image_dimensions
from .xml_helpers import (
    set_run_color, para_spacing,
    set_cell_bg, set_cell_borders, set_cell_margins, set_table_border,
    set_paragraph_border_bottom,
    set_row_cant_split, set_para_keep_next,
)
from .styles import apply_heading_style
from .metadata import HeadingNumberer


# Schemes that can become a live external hyperlink in a DOCX. A relationship
# Target is handed to the reader's OS to open, so the set stays deliberately
# small: no javascript:, no file:, no data:.
_LINKABLE_SCHEMES = ("http://", "https://", "mailto:", "ftp://", "ftps://")

# RFC 3986 characters that are already legal in a URI, plus '%' so existing
# percent-escapes survive. Without '%' in the safe set, a SharePoint URL
# containing %20 would re-encode to %2520 and stop resolving.
_URI_SAFE = "/:?#[]@!$&'()*+,;=~%"


def normalize_href(raw: str | None) -> str | None:
    """
    Return an absolute URL safe to use as a DOCX relationship target, or None
    when the reference should not become a live hyperlink.

    None is returned for relative paths ('../patterns/foo.md') and bare
    fragments ('#overview'). Relative links are resolved to absolute URLs
    earlier in the pipeline when the repo base URL is configured; anything
    still relative here has no absolute form, and a relationship pointing at
    it would resolve against wherever the reader happens to have saved the
    file. Styled-but-inert text is the honest rendering of that.

    Percent-encoding is applied only to characters that are illegal in a URI.
    Input that is already encoded passes through unchanged, so a SharePoint
    path ('.../Shared%20Documents/...') survives intact while a raw space
    becomes %20.
    """
    if not raw:
        return None

    url = raw.strip()
    if not url or url.startswith("#"):
        return None

    if not url.lower().startswith(_LINKABLE_SCHEMES):
        return None

    # quote() leaves unreserved characters alone; _URI_SAFE adds the reserved
    # set and '%'. Anything else - spaces, quotes, non-ASCII - is encoded as
    # UTF-8 percent-escapes, which is what Word expects in a Target.
    return quote(url, safe=_URI_SAFE, encoding="utf-8", errors="replace")


def open_hyperlink(para, href):
    """
    Append an empty w:hyperlink to *para* for *href* and return it, or return
    None when href is not linkable. Runs appended to the returned element
    become clickable; runs added to the paragraph directly do not.

    Shared by the body-text walker and the table-cell renderer, which are
    otherwise separate code paths - table cells never reach the HTML walker,
    which is why links in tables rendered as raw markdown until this existed.
    """
    target = normalize_href(href)
    if not target:
        return None
    try:
        r_id = para.part.relate_to(target, RT.HYPERLINK, is_external=True)
    except Exception:
        # A malformed target must not abort a 100-document build; the text
        # still renders, just without a live link.
        return None
    container = OxmlElement("w:hyperlink")
    container.set(qn("r:id"), r_id)
    para._p.append(container)
    return container


def _strip_html(text: str) -> str:
    """Remove HTML tags for plain text extraction."""
    return re.sub(r'<[^>]+>', '', text or '').strip()


# ── Emoji font splitter ───────────────────────────────────────────────────────
#
# Word cannot find emoji glyphs in Calibri or most body fonts.  When a run
# contains emoji, the characters must be emitted with a font that has coverage —
# on Windows, "Segoe UI Emoji" is the correct choice.
#
# _EMOJI_RE matches the Unicode ranges that cover the emoji used in this
# document set (colored circles, check/cross marks, misc symbols).
# _split_for_emoji splits a string into [(segment, is_emoji), ...] tuples so
# _add_run can emit separate runs with the appropriate font per segment.

_EMOJI_RE = re.compile(
    '['
    '\U0001F300-\U0001FAFF'   # Misc symbols, emoticons, supplemental (🔴🟡🟢✅❌ etc.)
    '\u2600-\u27BF'            # Misc symbols + dingbats (⚠️ ⚪ ☑ ☐ etc.)
    '\uFE00-\uFE0F'            # Variation selectors (emoji modifiers)
    ']+',
    re.UNICODE
)

def _split_for_emoji(text: str) -> list[tuple[str, bool]]:
    """Split *text* into [(segment, is_emoji), ...] pairs."""
    if not text:
        return []
    segments: list[tuple[str, bool]] = []
    prev = 0
    for m in _EMOJI_RE.finditer(text):
        if m.start() > prev:
            segments.append((text[prev:m.start()], False))
        segments.append((m.group(), True))
        prev = m.end()
    if prev < len(text):
        segments.append((text[prev:], False))
    return segments or [(text, False)]


# ── Inline markdown renderer for table cells ─────────────────────────────────
#
# Table cells are extracted before mistune sees them, so inline markdown
# (**bold**, *italic*, `code`) inside cells is never parsed by HtmlToDocx.
# These helpers tokenize inline markdown directly and emit styled runs.

_INLINE_MD_RE = re.compile(
    r'(`[^`\n]+`)'              # group 1 — code span
    r'|(\*\*\*[^*\n]+\*\*\*)'  # group 2 — bold + italic  ***...***
    r'|(\*\*[^*\n]+\*\*)'      # group 3 — bold  **...**
    r'|(__[^_\n]+__)'           # group 4 — bold  __...__
    r'|(\*[^*\n]+\*)'          # group 5 — italic  *...*
    r'|((?<!\w)_(?!_)[^_\n]+_(?!_)(?!\w))'  # group 6 — italic  _..._ (not __, not inside words)
    # Named so the numbered groups above keep their indices. Placed last is
    # safe: the alternatives above can only start at a backtick, asterisk, or
    # underscore, so at a '[' this is the only one that can match.
    #
    # The destination allows one level of balanced parentheses, as CommonMark
    # does. A plain [^)\s]* terminates at the first ')', which truncates a URL
    # like .../Function_(mathematics) mid-path and leaves a stray ')' in the
    # cell.
    r'|(?P<link>\[(?P<ltext>[^\]\n]*)\]'
    r'\((?P<lurl>(?:[^()\s]|\([^()\s]*\))*)\))'  # [text](url)
    # Autolink <https://...>. Without this a cell containing one renders the
    # angle brackets literally, which is what a reader saw in the external
    # references table of a real document. Restricted to a scheme-looking
    # target so ordinary angle-bracketed prose in a cell - <placeholder>,
    # <name> - is left alone rather than being swallowed as a link.
    r'|(?P<auto><(?P<aurl>[A-Za-z][A-Za-z0-9+.-]*:[^>\s]+)>)'
)


def _add_cell_run(para, text, *, bold, italic, code, color, font_name, font_size,
                  underline=False, container=None):
    """Add a single formatted run to a table-cell paragraph.

    Emoji characters are split into separate runs with Segoe UI Emoji font so
    Word can locate the glyphs.  Code spans are never split (no emoji expected).

    *container* is an open w:hyperlink element from open_hyperlink(); when
    given, runs are moved inside it so the whole span is clickable.
    """
    if not text:
        return
    segments = [(text, False)] if code else _split_for_emoji(text)
    for segment_text, is_emoji in segments:
        if not segment_text:
            continue
        run           = para.add_run(segment_text)
        run.bold      = bold
        run.italic    = italic
        run.font.name = ("Segoe UI Emoji" if is_emoji and not code else font_name)
        run.font.size = font_size
        run.underline = underline
        set_run_color(run, color)
        if container is not None:
            container.append(run._element)


def _render_cell_text(para, text, *, base_bold=False, color, font_name, font_size):
    """
    Parse inline markdown in *text* and add styled runs to *para*.

    Recognised constructs:
      ``code``             → monospace, dark-gray, one point below the cell
      ***bold+italic***    → bold + italic
      **bold** / __bold__  → bold
      *italic* / _italic_  → italic

    Anything not matched is emitted as a plain run with *base_bold* and
    *color*. This function is used for both header cells (base_bold=True,
    white color) and body cells (base_bold=False, black color).
    """
    pos = 0
    for m in _INLINE_MD_RE.finditer(text):
        # Plain text before this match
        if m.start() > pos:
            _add_cell_run(para, text[pos:m.start()],
                          bold=base_bold, italic=False, code=False,
                          color=color, font_name=font_name, font_size=font_size)

        matched = m.group(0)
        if m.group('link'):                     # [text](url)
            label = m.group('ltext') or m.group('lurl')
            container = open_hyperlink(para, m.group('lurl'))
            # A relative or fragment target yields no container. The label
            # still renders in link styling - the reader sees the same words
            # they would in GitHub, just without a destination - which beats
            # showing them raw [text](path) markdown.
            _add_cell_run(para, label,
                          bold=base_bold, italic=False, code=False,
                          color=BLUE_LINK, font_name=font_name,
                          font_size=font_size, underline=True,
                          container=container)
        elif m.group('auto'):                   # <https://example.com>
            # The URL is both the label and the target; the brackets are
            # syntax and never render.
            url = m.group('aurl')
            container = open_hyperlink(para, url)
            _add_cell_run(para, url,
                          bold=base_bold, italic=False, code=False,
                          color=BLUE_LINK, font_name=font_name,
                          font_size=font_size, underline=True,
                          container=container)
        elif m.group(1):                        # `code`
            _add_cell_run(para, matched[1:-1],
                          bold=False, italic=False, code=True,
                          color=RGBColor(0x1F, 0x1F, 0x1F),
                          font_name=FONT_MONO,
                          font_size=Pt(SIZE_CODE_IN_CELL))
        elif m.group(2):                        # ***bold+italic***
            _add_cell_run(para, matched[3:-3],
                          bold=True, italic=True, code=False,
                          color=color, font_name=font_name, font_size=font_size)
        elif m.group(3) or m.group(4):         # **bold** or __bold__
            _add_cell_run(para, matched[2:-2],
                          bold=True, italic=False, code=False,
                          color=color, font_name=font_name, font_size=font_size)
        elif m.group(5) or m.group(6):         # *italic* or _italic_
            _add_cell_run(para, matched[1:-1],
                          bold=base_bold, italic=True, code=False,
                          color=color, font_name=font_name, font_size=font_size)

        pos = m.end()

    # Trailing plain text
    if pos < len(text):
        _add_cell_run(para, text[pos:],
                      bold=base_bold, italic=False, code=False,
                      color=color, font_name=font_name, font_size=font_size)


# ── HTML → docx walker ────────────────────────────────────────────────────────

class HtmlToDocx(HTMLParser):
    """
    Walk mistune-rendered HTML and emit python-docx paragraphs into `doc`.

    Headings are collected in self.headings as (level, text, None) tuples
    for use by the TOC builder.

    State tracking:
      _tag_stack     : stack of (tag, attrs_dict) for active open tags
      _list_stack    : stack of ('ul'|'ol', [counter]) for nested lists
      _current_para  : the paragraph currently being built
      _in_pre        : True while inside a <pre> block
      _pre_buf       : accumulates raw text inside <pre>
      _in_blockquote : True while inside <blockquote>
    """

    def __init__(self, doc, md_src_dir: str | None = None):
        super().__init__()
        self.doc        = doc
        self.headings   = []   # (level, numbered_text, None)
        # Directory of the source .md file, used to resolve relative image paths.
        # Defaults to the current working directory if not provided.
        self.md_src_dir = md_src_dir or os.getcwd()

        self._current_para   = None
        self._tag_stack      = []
        self._list_stack     = []
        self._in_pre         = False
        self._pre_buf        = []
        self._in_blockquote  = False
        self._skip_tags      = {'html', 'body', 'head'}
        self._heading_numberer = HeadingNumberer()

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _current_tags(self):
        return {t for t, _ in self._tag_stack}

    def _current_href(self):
        """href of the innermost open <a>, or None. Searches from the top of
        the stack so a nested anchor - which mistune will not emit, but a
        hand-written HTML fragment could - resolves to the nearest one."""
        for tag, attrs in reversed(self._tag_stack):
            if tag == 'a':
                return attrs.get('href')
        return None

    def _flush_para(self):
        self._current_para = None

    def _ensure_para(self, before=60, after=80):
        if self._current_para is None:
            self._current_para = self.doc.add_paragraph()
            para_spacing(self._current_para, before=before, after=after)
        return self._current_para

    def _add_run(self, text, bold=False, italic=False, code=False, link=False,
                 href=None):
        if not text:
            return
        para = self._ensure_para()

        # A live hyperlink is a w:hyperlink element carrying a relationship id,
        # wrapping the runs. Link styling alone (blue + underline) only looks
        # clickable; without the relationship there is nothing to click.
        container = open_hyperlink(para, href) if link else None

        # Split on emoji boundaries so emoji segments get Segoe UI Emoji font.
        # Code spans are never emoji; skip splitting for them.
        segments = [(text, False)] if code else _split_for_emoji(text)
        for segment_text, is_emoji in segments:
            if not segment_text:
                continue
            run           = para.add_run(segment_text)
            run.font.name = (FONT_MONO if code
                             else "Segoe UI Emoji" if is_emoji else FONT_BODY)
            run.font.size = Pt(SIZE_CODE if code else SIZE_BODY)
            run.bold      = bold
            run.italic    = italic
            if link:
                set_run_color(run, BLUE_LINK)
                run.underline = True
            elif code:
                set_run_color(run, RGBColor(0x1F, 0x1F, 0x1F))
            else:
                set_run_color(run, BLACK)
            # add_run appended to the paragraph; move it inside the hyperlink
            # so the whole link text is clickable, emoji segments included.
            if container is not None:
                container.append(run._element)

    def _handle_img(self, attrs_dict: dict) -> None:
        """
        Embed an image referenced by an <img src="..." alt="..."> tag.

        src is resolved relative to self.md_src_dir so that relative paths
        in the markdown (e.g. ../../diagrams/exported/foo.png) work correctly.
        alt text is rendered as a small italic caption below the image.

        Missing or unresolvable paths emit a styled warning paragraph instead
        of raising an exception, so the build always completes.

        Supported formats: PNG, JPG/JPEG, GIF, TIFF, BMP.
        SVG is not supported by python-docx — export to PNG first.
        """
        src = attrs_dict.get('src', '').strip()
        alt = attrs_dict.get('alt', '').strip()

        if not src:
            return

        # Resolve relative paths from the source .md file's directory
        img_path = src if os.path.isabs(src) else \
            os.path.normpath(os.path.join(self.md_src_dir, src))

        self._flush_para()

        if os.path.isfile(img_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para_spacing(p, before=120, after=40)
            width, height = fit_image_dimensions(img_path)
            if height is None:
                p.add_run().add_picture(img_path, width=width)
            else:
                p.add_run().add_picture(img_path, width=width, height=height)
            if alt:
                cap = self.doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para_spacing(cap, before=0, after=120)
                run           = cap.add_run(alt)
                run.italic    = True
                run.font.name = FONT_BODY
                run.font.size = Pt(SIZE_CAPTION)
                set_run_color(run, GRAY_TEXT)
        else:
            # Image not found — warn without crashing
            p   = self.doc.add_paragraph()
            para_spacing(p, before=60, after=60)
            run = p.add_run(f'[Image not found: {src}]')
            run.italic    = True
            run.font.name = FONT_BODY
            run.font.size = Pt(SIZE_BODY)
            set_run_color(run, RGBColor(0xCC, 0x00, 0x00))

    # ── HTMLParser callbacks ──────────────────────────────────────────────────

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)

        # img is a void element — handle fully here without touching the tag
        # stack.  HTMLParser may call handle_starttag + handle_endtag for
        # self-closing <img/>, but since we never push img onto _tag_stack
        # the subsequent handle_endtag is harmless (stack check will miss).
        if tag == 'img':
            self._handle_img(attrs_dict)
            return

        self._tag_stack.append((tag, attrs_dict))

        if tag in self._skip_tags:
            return

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            self._flush_para()
            self._current_para = self.doc.add_paragraph()

        elif tag == 'p':
            self._flush_para()
            if self._in_blockquote:
                self._current_para = self.doc.add_paragraph()
                para_spacing(self._current_para, before=60, after=60)
                self._current_para.paragraph_format.left_indent = Inches(0.5)
            else:
                self._current_para = self.doc.add_paragraph()
                para_spacing(self._current_para, before=60, after=80)

        elif tag in ('ul', 'ol'):
            self._list_stack.append((tag, [0]))

        elif tag == 'li':
            self._flush_para()
            depth   = len(self._list_stack)
            ordered = self._list_stack[-1][0] == 'ol' if self._list_stack else False
            if ordered:
                self._list_stack[-1][1][0] += 1
                prefix = f"{self._list_stack[-1][1][0]}."
            else:
                prefix = "\u2022"   # bullet

            indent_left_in    = (360 * depth) / 1440
            indent_hanging_in = 360 / 1440

            self._current_para = self.doc.add_paragraph()
            para_spacing(self._current_para, before=40, after=40)
            self._current_para.paragraph_format.left_indent      = Inches(indent_left_in + indent_hanging_in)
            self._current_para.paragraph_format.first_line_indent = Inches(-indent_hanging_in)
            run = self._current_para.add_run(f"{prefix}\t")
            run.font.name = FONT_BODY
            run.font.size = Pt(SIZE_BODY)

        elif tag == 'blockquote':
            self._in_blockquote = True

        elif tag == 'pre':
            self._in_pre  = True
            self._pre_buf = []

        elif tag == 'hr':
            p = self.doc.add_paragraph()
            set_paragraph_border_bottom(p, color="CCCCCC", sz=4)
            para_spacing(p, before=120, after=120)
            self._flush_para()

    def handle_endtag(self, tag):
        if self._tag_stack and self._tag_stack[-1][0] == tag:
            self._tag_stack.pop()

        if tag in ('h1', 'h2', 'h3', 'h4', 'h5', 'h6'):
            if self._current_para:
                level    = min(int(tag[1]), 3)   # clamp h4–h6 to h3 style
                # Walk every w:t descendant rather than paragraph.runs.
                # python-docx's .runs returns only direct w:r children, so a
                # run moved inside a w:hyperlink is invisible to it - which
                # silently dropped the link text out of a heading, and left a
                # heading that was entirely a link completely blank.
                raw_text = ''.join(
                    node.text or ''
                    for node in self._current_para._p.iter(qn('w:t'))
                )
                numbered = self._heading_numberer.format(level, raw_text)
                # Replace paragraph content with the numbered heading text
                # so the heading in the document body and the TOC field match.
                self._current_para.clear()
                self._current_para.add_run(numbered)
                apply_heading_style(self.doc, self._current_para, level)
                self.headings.append((level, numbered, None))
            self._flush_para()

        elif tag == 'p':
            self._flush_para()

        elif tag in ('ul', 'ol'):
            if self._list_stack:
                self._list_stack.pop()
            self._flush_para()

        elif tag == 'li':
            self._flush_para()

        elif tag == 'blockquote':
            self._in_blockquote = False
            self._flush_para()

        elif tag == 'pre':
            self._in_pre = False
            text = ''.join(self._pre_buf)
            p    = self.doc.add_paragraph()
            para_spacing(p, before=60, after=60)
            # Light gray background for code blocks via raw XML shading
            from docx.oxml import OxmlElement as _OxmlElement
            from docx.oxml.ns import qn as _qn
            pPr = p._p.get_or_add_pPr()
            shd = _OxmlElement('w:shd')
            shd.set(_qn('w:val'),   'clear')
            shd.set(_qn('w:color'), 'auto')
            shd.set(_qn('w:fill'),  'F2F2F2')
            pPr.append(shd)
            run = p.add_run(text)
            run.font.name = FONT_MONO
            run.font.size = Pt(SIZE_CODE)
            set_run_color(run, RGBColor(0x1F, 0x1F, 0x1F))
            self._flush_para()

    def handle_data(self, data):
        if not data:
            return

        if self._in_pre:
            self._pre_buf.append(data)
            return

        # Don't create a new paragraph just for inter-element whitespace when
        # no paragraph is open (e.g. whitespace inside <p>...</p> after an
        # embedded <img> has already flushed _current_para).
        if data.strip() == '' and self._current_para is None:
            return

        tags = self._current_tags()

        # Ignore whitespace-only text between block-level elements
        if data.strip() == '' and not (tags & {'p', 'li', 'h1', 'h2', 'h3',
                                               'h4', 'h5', 'h6',
                                               'strong', 'em', 'code', 'a',
                                               'blockquote'}):
            return

        bold   = 'strong' in tags
        italic = 'em'     in tags or self._in_blockquote
        code   = 'code'   in tags and 'pre' not in tags
        link   = 'a'      in tags

        self._add_run(data, bold=bold, italic=italic, code=code, link=link,
                      href=self._current_href())


# ── GFM table extraction ──────────────────────────────────────────────────────

def extract_md_tables(md_text: str) -> tuple[str, list[dict]]:
    """
    Extract GFM pipe tables from markdown without touching the document.

    Returns (processed_text, table_data_list) where:
      - processed_text has each table replaced by a __TABLE_N__ placeholder
      - table_data_list[N] is a dict with all data needed to render table N

    Call render_md_table(doc, table_data) for each placeholder encountered
    while iterating segments in order. This keeps tables and body text
    interleaved correctly in the final document.

    Table styling applied by render_md_table:
      - Header row: BLUE_MID (#2E75B6) background, white text
      - Body rows:  alternating white / light gray (#F2F2F2)
      - Column widths: evenly distributed across 8" content width
      - Alignment: derived from GFM divider row (:---:, ---:, :---)
    """
    table_re = re.compile(
        r'(?m)^(\|.+\|\n)(\|[-| :]+\|\n)((?:\|.+\|\n?)*)',
        re.MULTILINE
    )
    tables: list[dict] = []

    def extract_table(m):
        header_row  = [c.strip() for c in m.group(1).strip().split('|') if c.strip()]
        divider_row = [c.strip() for c in m.group(2).strip().split('|') if c.strip()]
        body_rows   = []
        for line in m.group(3).strip().splitlines():
            row = [c.strip() for c in line.strip().split('|') if c.strip()]
            if row:
                body_rows.append(row)

        idx = len(tables)
        tables.append({
            'header_row':  header_row,
            'divider_row': divider_row,
            'body_rows':   body_rows,
        })
        return f'\n__TABLE_{idx}__\n'

    processed = table_re.sub(extract_table, md_text)
    return processed, tables


_MD_INLINE_RE = re.compile(
    r'\[([^\]]*)\]\([^)]*\)'      # [text](url) -> text
    r'|<([^>\s]+)>'               # <url>       -> url
    r'|[`*_]+'                    # emphasis and code marks
)


def _visible_len(cell: str) -> str:
    """Cell text as the reader sees it, with inline markdown syntax removed.

    Width should follow rendered text, not source. Without this a cell of
    ``**Recommended target**`` claims four characters it never draws.
    """
    return _MD_INLINE_RE.sub(lambda m: m.group(1) or m.group(2) or '', cell)


def _column_widths(header_row, body_rows, n_cols,
                   total_in=TABLE_TOTAL_WIDTH_IN,
                   min_in=TABLE_MIN_COL_IN,
                   char_cap=TABLE_COL_CHAR_CAP):
    """Distribute table width across columns by how much text each holds.

    Even distribution is what makes rendered tables look wrong: a column of
    single digits claims the same inches as a column of sentences, so one
    side is a field of whitespace while the other wraps every row.

    Each column asks for width on two grounds:

    * its **longest word**, because a word that does not fit is a word that
      breaks badly or overflows the cell; and
    * its **longest cell**, capped at *char_cap* — past that point a cell is
      going to wrap no matter what it gets, so letting it bid its full length
      would starve every other column to no benefit.

    Both bids are capped. The word bid needs it as much as the cell bid does:
    a long URL is a single unbreakable token, and uncapped it would ask for
    its entire length and leave the rest of the table nothing. A token longer
    than the cap gets broken by the renderer whatever width it is given, so
    bidding past the cap buys damage rather than avoiding it.

    Widths are then allocated in proportion to those bids, and any column
    landing under *min_in* is raised to it, with the shortfall taken from the
    columns that have room to give. A single-digit column ends up narrow but
    still legible, which is the whole point.
    """
    def bid(idx):
        texts = [_visible_len(header_row[idx] if idx < len(header_row) else '')]
        texts += [_visible_len(row[idx]) for row in body_rows if idx < len(row)]
        longest_word = max((len(w) for t in texts for w in t.split()), default=1)
        longest_cell = max((len(t) for t in texts), default=1)
        return float(max(min(longest_word, char_cap),
                         min(longest_cell, char_cap), 1))

    bids  = [bid(i) for i in range(n_cols)]
    total_bid = sum(bids) or 1.0
    widths = [total_in * b / total_bid for b in bids]

    # Raise anything below the floor, and fund it from the columns above the
    # floor in proportion to their surplus. Guard the degenerate case where
    # every column is at or under the floor (a very wide table of tiny cells):
    # there is nothing to take from, so leave the even split alone.
    deficit = sum(min_in - w for w in widths if w < min_in)
    if deficit > 0:
        surplus = sum(w - min_in for w in widths if w > min_in)
        if surplus > deficit:
            widths = [
                min_in if w < min_in
                else w - (w - min_in) * (deficit / surplus)
                for w in widths
            ]
        else:
            widths = [total_in / n_cols] * n_cols
    return widths


def render_md_table(doc, table_data: dict) -> None:
    """
    Render a single parsed GFM table into the document at the current position.

    Called from builder.py while iterating body segments in order so that
    each table lands immediately after the preceding paragraph — not all
    tables first followed by all text (the old broken behaviour).
    """
    header_row  = table_data['header_row']
    divider_row = table_data['divider_row']
    body_rows   = table_data['body_rows']

    # Derive alignment from divider cells
    alignments = []
    for cell in divider_row:
        if cell.startswith(':') and cell.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif cell.endswith(':'):
            alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            alignments.append(WD_ALIGN_PARAGRAPH.LEFT)

    n_cols = len(header_row)
    col_widths = _column_widths(header_row, body_rows, n_cols)

    table = doc.add_table(rows=1, cols=n_cols)
    table.style     = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_border(table, color="CCCCCC", sz=4)

    # Header row
    hdr = table.rows[0]
    for i, col_text in enumerate(header_row):
        cell = hdr.cells[i]
        cell.width = Inches(col_widths[i])
        set_cell_bg(cell, "2E75B6")
        set_cell_borders(cell, color="CCCCCC", sz=4)
        set_cell_margins(cell)
        p   = cell.paragraphs[0]
        p.alignment = alignments[i] if i < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
        para_spacing(p, before=36, after=36)
        _render_cell_text(p, col_text,
                          base_bold=True,
                          color=RGBColor(0xFF, 0xFF, 0xFF),
                          font_name=FONT_BODY,
                          font_size=Pt(SIZE_TABLE_HEADER))

    # Body rows
    for ridx, body_row in enumerate(body_rows):
        row  = table.add_row()
        fill = "F2F2F2" if ridx % 2 == 1 else "FFFFFF"
        for i in range(n_cols):
            cell = row.cells[i]
            cell.width = Inches(col_widths[i])
            set_cell_bg(cell, fill)
            set_cell_borders(cell, color="CCCCCC", sz=4)
            set_cell_margins(cell)
            p   = cell.paragraphs[0]
            p.alignment = alignments[i] if i < len(alignments) else WD_ALIGN_PARAGRAPH.LEFT
            para_spacing(p, before=36, after=36)
            val = body_row[i] if i < len(body_row) else ''
            _render_cell_text(p, val,
                              base_bold=False,
                              color=BLACK,
                              font_name=FONT_BODY,
                              font_size=Pt(SIZE_TABLE_CELL))

    # Keep the table together on one page unless it is too large to fit.
    #
    # Strategy:
    #   1. cantSplit on every row — prevents any row from being split by a
    #      page break on its own.
    #   2. keepNext on all cell paragraphs except the very last cell — this
    #      chains every row to the next through Word's keep-with-next layout
    #      pass, causing Word to push the whole table to the next page if it
    #      does not fit in the remaining space. If the table exceeds a full
    #      page, Word breaks it normally (keepNext is advisory, not absolute).
    all_rows   = table.rows
    n_rows     = len(all_rows)
    for r_idx, row in enumerate(all_rows):
        set_row_cant_split(row)
        is_last_row = (r_idx == n_rows - 1)
        for c_idx, cell in enumerate(row.cells):
            is_last_cell = is_last_row and (c_idx == n_cols - 1)
            if not is_last_cell:
                for para in cell.paragraphs:
                    set_para_keep_next(para)

    doc.add_paragraph()   # spacer after table
